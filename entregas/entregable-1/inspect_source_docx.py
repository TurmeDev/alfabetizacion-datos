import argparse
import os
import sys
from pathlib import Path

from docx import Document


SCRIPT_DIR = Path(__file__).resolve().parent
DEFAULT_SOURCE = SCRIPT_DIR / "Guia_Entregable_1_Startup_Educativa.docx"


def clean(text: str) -> str:
    return " / ".join(part.strip() for part in text.splitlines() if part.strip())


def main() -> None:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")

    parser = argparse.ArgumentParser(description="Inspecciona la estructura del DOCX fuente.")
    parser.add_argument(
        "source",
        nargs="?",
        type=Path,
        default=Path(os.environ.get("ENTREGABLE_SOURCE", str(DEFAULT_SOURCE))),
        help="DOCX fuente; también puede definirse con ENTREGABLE_SOURCE.",
    )
    args = parser.parse_args()
    source = args.source.expanduser().resolve()
    if not source.is_file():
        parser.error(
            f"No se encontró el DOCX fuente: {source}. "
            "Indicá la ruta o definí ENTREGABLE_SOURCE."
        )

    print("exists", True, "size", source.stat().st_size)
    doc = Document(str(source))
    print("paragraphs", len(doc.paragraphs), "tables", len(doc.tables), "sections", len(doc.sections))
    for i, paragraph in enumerate(doc.paragraphs):
        text = clean(paragraph.text)
        if text:
            print(f"P{i:03d} [{paragraph.style.name}] {text[:240]}")
    print("--- TABLES ---")
    for ti, table in enumerate(doc.tables):
        print(f"TABLE {ti}: rows={len(table.rows)} cols={len(table.columns)}")
        for ri, row in enumerate(table.rows[:20]):
            cells = [clean(cell.text)[:120] for cell in row.cells]
            print(f"  R{ri:02d}: {cells}")


if __name__ == "__main__":
    main()
