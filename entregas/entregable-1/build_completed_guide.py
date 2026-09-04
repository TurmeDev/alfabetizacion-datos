from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


SOURCE = Path(r"C:\Users\User\Downloads\Guia_Entregable_1_Startup_Educativa.docx")
OUT_DIR = Path("entregas/entregable-1/generated")
OUTPUT = OUT_DIR / "Guia_Entregable_1_Startup_Educativa_COMPLETADA_CriterIA.docx"

FLOW_IMAGES = [
    OUT_DIR / "flujograma-01_inicio_actividad_1.png",
    OUT_DIR / "flujograma-02_actividad_2_3.png",
    OUT_DIR / "flujograma-03_actividad_4.png",
    OUT_DIR / "flujograma-04_actividad_5.png",
    OUT_DIR / "flujograma-05_actividad_6.png",
    OUT_DIR / "flujograma-06_actividad_7.png",
    OUT_DIR / "flujograma-07_actividad_8.png",
    OUT_DIR / "flujograma-08_actividad_9_fin.png",
]


DESCRIPTION = (
    '"Dos Tonos" es una experiencia de alfabetización de datos para estudiantes '
    "de educación media dentro de la Estación 2 del proyecto Startup Educativa. "
    "El estudiante realiza una consulta sobre candidatos políticos y observa dos "
    "respuestas generadas a partir de la misma base de información verificada. "
    "Una respuesta se presenta en tono neutro y otra con un encuadre "
    "intencionalmente diferente, para que el estudiante compare cómo cambia la "
    "percepción cuando se resaltan u omiten ciertos aspectos. La actividad no "
    "busca orientar el voto, sino enseñar que los datos, aunque sean los mismos, "
    "pueden comunicarse de maneras distintas. Durante la estación, el equipo guía "
    "la comparación, registra observaciones acordadas con Control de Gestión y "
    "deja evidencia del proceso para su evaluación posterior."
)


ACTIVITIES = [
    [
        "1",
        "Analizar la guía y el contexto",
        "Todos",
        "Reunión inicial",
        "Guía, mapa visual y consignas del docente",
        "Guía comprendida y dudas listadas",
    ],
    [
        "2",
        "Diseñar el servicio",
        "Todos",
        "Sesión de diseño",
        "Notas de la actividad 1",
        "Dinámica educativa definida y explicable",
    ],
    [
        "3",
        "Seleccionar la herramienta de apoyo",
        "Lucas y Diego",
        "Trabajo individual y puesta en común",
        "Acceso a herramientas de IA",
        "Herramienta elegida para la estación",
    ],
    [
        "4",
        "Preparar fichas de candidatos ficticios",
        "Arnold y Mathias",
        "Trabajo individual y revisión cruzada",
        "ADR-0003, plantilla y checklist",
        "Tres fichas sintéticas aprobadas",
    ],
    [
        "5",
        'Preparar el chatbot "Dos Tonos"',
        "Lucas y Mathias",
        "Trabajo de preparación",
        "Base de candidatos y herramienta seleccionada",
        "Chatbot listo para demostración",
    ],
    [
        "6",
        "Preparar la demostración de la Estación 2",
        "Todos + sub-equipo B",
        "Trabajo previo al evento",
        "Guion y lista de verificación",
        "Estación lista para operar",
    ],
    [
        "7",
        "Prestar el servicio en el colegio",
        "Todos + sub-equipo B",
        "Colegio asignado",
        "Chatbot, dispositivo e internet",
        "Experiencia realizada y datos registrados",
    ],
    [
        "8",
        "Cerrar la estación y ordenar evidencia",
        "Mathias y Arnold + Módulo 5",
        "Cierre posterior",
        "Observaciones, incidencias y datos",
        "Evidencia ordenada y usable",
    ],
    [
        "9",
        "Elaborar el informe académico",
        "CriterIA completo",
        "Trabajo final",
        "Evidencia del servicio y guía",
        "Documento completo y alineado al servicio",
    ],
]


INTERVENIENTES = [
    [
        "Equipo CriterIA (sub-equipo A, Módulo 2)",
        "Diseña el servicio, prepara el chatbot, valida la estación, acompaña la demostración y elabora el informe.",
    ],
    [
        "Sub-equipo B de Módulo 2",
        "Apoya la preparación de contenidos, materiales y facilitación durante la Estación 2.",
    ],
    [
        "Módulo 3: Logística y operación",
        "Coordina condiciones del evento: espacios, tiempos, movilidad, recepción y soporte operativo.",
    ],
    [
        "Módulo 5: Evaluación y análisis",
        "Define y recibe los datos/evidencias necesarios para analizar resultados e impacto.",
    ],
    [
        "Estudiantes de educación media",
        "Participan en la experiencia, realizan consultas y comparan las dos respuestas del chatbot.",
    ],
]


RECURSOS = [
    ["Guía del Entregable 1", "1, 2 y 9", "Equipo CriterIA"],
    ["Mapa visual del proyecto", "1, 2, 6 y 7", "CriterIA, sub-equipo B y Módulo 3"],
    ["Herramienta de IA / chatbot", "3, 5, 6 y 7", "Equipo CriterIA"],
    ["ADR-0003, plantilla y fichas sintéticas", "4 y 5", "Arnold y Mathias"],
    ["Base documental de candidatos", "4, 5 y 7", "Equipo CriterIA"],
    ["Guion de demostración y lista de verificación", "6 y 7", "CriterIA y sub-equipo B"],
    ["Dispositivo e internet", "6 y 7", "CriterIA y Módulo 3"],
    ["Registro de observaciones, incidencias y datos", "7 y 8", "CriterIA y Módulo 5"],
]


CONTROLES = [
    [
        "Comprensión inicial del encargo",
        "Si la guía, alcance y rol del equipo están claros.",
        "Equipo CriterIA",
        "Si hay dudas, se revisa la guía/mapa antes de diseñar.",
    ],
    [
        "Diseño del servicio",
        "Si la dinámica explica alfabetización de datos sin tecnicismo excesivo.",
        "Equipo CriterIA",
        "Si no es claro, se ajusta el objetivo educativo y el guion.",
    ],
    [
        "Herramienta de apoyo",
        "Si la herramienta permite operar el chatbot en la estación.",
        "Delegados + CriterIA",
        "Si no es viable, se elige otra opción disponible.",
    ],
    [
        "Base de candidatos",
        "Si las fichas están completas y verificadas.",
        "Delegados + revisión cruzada",
        "Si falta información, se corrige antes de preparar el chatbot.",
    ],
    [
        "Validación del chatbot",
        "Si ambos tonos leen la misma base y la diferencia está en la presentación.",
        "Equipo CriterIA",
        "Si falla, se corrigen instrucciones, ejemplos o fichas.",
    ],
    [
        "Registro para evaluación",
        "Si se guardaron observaciones y datos acordados con Módulo 5.",
        "CriterIA + Módulo 5",
        "Si falta evidencia, se registra la incidencia y se completa lo posible.",
    ],
]


CONCLUSION = (
    'Al analizar "Dos Tonos" como un proceso de servicio, el equipo CriterIA '
    "aprendió tres cosas centrales. Primero, descomponer la Estación 2 en "
    "actividades discretas hace visibles dependencias que podrían pasar "
    "desapercibidas: la demostración depende de contar con una herramienta "
    "operativa, una base de información verificada, apoyo de facilitación, "
    "condiciones logísticas y datos definidos para evaluación. Segundo, los "
    "puntos de control son tan importantes como las actividades mismas: cada "
    "avance debe cerrar con una verificación concreta, como guía comprendida, "
    "dinámica educativa confirmada, fichas verificadas, chatbot validado, "
    "estación lista, datos registrados y evidencia ordenada. Tercero, el "
    "servicio debe explicar la IA sin tecnicismo excesivo: el aprendizaje central "
    "no es programar un modelo, sino observar que una misma base de datos puede "
    "presentarse con distintos tonos y producir percepciones diferentes.\n\n"
    "Los principales aspectos a mejorar antes de la implementación son: definir "
    "con precisión el alcance del sub-equipo B, construir la base de candidatos, "
    "confirmar la herramienta de apoyo, probar el chatbot en condiciones reales "
    "y acordar con Control de Gestión qué datos se registrarán. El flujograma "
    "ayuda a distinguir qué puede prepararse en paralelo y qué depende de una "
    "secuencia obligatoria."
)


def set_cell_text(cell, text: str, size: int = 9) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(size)


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = text
    for run in paragraph.runs:
        run.font.name = "Calibri"


def fill_tables(doc: Document) -> None:
    # Activities.
    table = doc.tables[0]
    for row_idx, activity in enumerate(ACTIVITIES, start=1):
        for col_idx, value in enumerate(activity):
            set_cell_text(table.cell(row_idx, col_idx), value, size=8)
    for row_idx in range(1 + len(ACTIVITIES), len(table.rows)):
        for cell in table.rows[row_idx].cells:
            set_cell_text(cell, "", size=8)

    # Intervenientes.
    table = doc.tables[1]
    while len(table.rows) < len(INTERVENIENTES) + 1:
        table.add_row()
    for row_idx, row in enumerate(INTERVENIENTES, start=1):
        for col_idx, value in enumerate(row):
            set_cell_text(table.cell(row_idx, col_idx), value, size=9)

    # Recursos.
    table = doc.tables[2]
    while len(table.rows) < len(RECURSOS) + 1:
        table.add_row()
    for row_idx, row in enumerate(RECURSOS, start=1):
        for col_idx, value in enumerate(row):
            set_cell_text(table.cell(row_idx, col_idx), value, size=8)

    # Decisiones/controles.
    table = doc.tables[3]
    while len(table.rows) < len(CONTROLES) + 1:
        table.add_row()
    for row_idx, row in enumerate(CONTROLES, start=1):
        for col_idx, value in enumerate(row):
            set_cell_text(table.cell(row_idx, col_idx), value, size=8)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    for image in FLOW_IMAGES:
        if not image.exists():
            raise FileNotFoundError(image)

    shutil.copy2(SOURCE, OUTPUT)
    doc = Document(str(OUTPUT))

    flow_was_filled = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Equipo:"):
            set_paragraph_text(
                paragraph,
                "Equipo: CriterIA — Sub-equipo A del Módulo 2 (Contenido y Demostraciones).",
            )
        elif text.startswith("Tema:"):
            set_paragraph_text(
                paragraph,
                "Tema: Módulos 1 y 2 / Estación 2: Alfabetización de Datos — chatbot y base documental.",
            )
        elif text.startswith("Nombre:"):
            set_paragraph_text(paragraph, 'Nombre: "Dos Tonos"')
        elif text.startswith("________________________________________________________________") and not flow_was_filled:
            set_paragraph_text(paragraph, DESCRIPTION)
        elif text.startswith("Nombre del proceso:"):
            set_paragraph_text(
                paragraph,
                "Nombre del proceso: Diseño, preparación y prestación de la Estación 2: Alfabetización de Datos.",
            )
        elif text.startswith("Inicio del proceso:"):
            set_paragraph_text(
                paragraph,
                "Inicio del proceso: el equipo CriterIA recibe el componente asignado y define el objetivo educativo de la estación.",
            )
        elif text.startswith("Fin del proceso:"):
            set_paragraph_text(
                paragraph,
                "Fin del proceso: estudiantes de educación media completan la experiencia y se registran los datos acordados para evaluación.",
            )
        elif text.startswith("Usuario/beneficiario:"):
            set_paragraph_text(
                paragraph,
                "Usuario/beneficiario: estudiantes de educación media que recorren la Estación 2.",
            )
        elif text.startswith("Resultado esperado:"):
            set_paragraph_text(
                paragraph,
                "Resultado esperado: experiencia educativa clara, guiada y medible sobre distintas formas de presentar una misma información.",
            )
        elif text == "[Insertar aquí el diagrama]":
            flow_was_filled = True
            set_paragraph_text(
                paragraph,
                "Diagrama preliminar del proceso: INICIO → analizar guía/contexto → diseñar servicio → seleccionar herramienta → recopilar fichas → preparar chatbot → preparar demostración → prestar servicio → ordenar evidencia → elaborar informe → FIN. Controles principales: ¿guía comprendida? → ¿objetivo claro? → ¿herramienta apta? → ¿información verificada? → ¿chatbot validado? → ¿estación lista? → ¿datos registrados? → ¿evidencia usable?",
            )
        elif text.startswith("Tipo de flujograma utilizado:"):
            set_paragraph_text(
                paragraph,
                "Tipo de flujograma utilizado:   ☒ Vertical    ☐ Horizontal    ☐ Otro: segmentado para lectura en Word/PDF",
            )
        elif text.startswith("Flujograma del proceso:"):
            set_paragraph_text(
                paragraph,
                "Flujograma del proceso: ver Anexo 1 al final del documento. Se presenta en segmentos para mantener la lectura. La secuencia completa corresponde a un flujograma vertical con decisiones y controles.",
            )
        elif text.startswith("Justificación de la elección:"):
            set_paragraph_text(
                paragraph,
                "Justificación de la elección: se utiliza un flujograma vertical porque permite leer el proceso de arriba hacia abajo, siguiendo la secuencia natural de preparación, prestación y cierre del servicio. La simbología distingue actividades, decisiones/controles y resultados. Para evitar que el diagrama formal quede ilegible en una sola página, se presenta segmentado en el documento.",
            )
        elif text.startswith("________________________________________________________________") and flow_was_filled:
            set_paragraph_text(paragraph, CONCLUSION)
        elif text.startswith("Nombre del archivo:"):
            set_paragraph_text(
                paragraph,
                "Nombre del archivo sugerido: Equipo_CriterIA_Entregable_1_StartupEducativa.pdf",
            )

    fill_tables(doc)

    doc.add_page_break()
    doc.add_heading("Anexo 1. Flujograma formal del proceso", level=1)
    note = doc.add_paragraph(
        "El flujograma se divide en segmentos para que pueda leerse en Word/PDF sin reducirlo hasta volverlo ilegible."
    )
    for run in note.runs:
        run.font.size = Pt(10)
    for idx, image in enumerate(FLOW_IMAGES, start=1):
        if idx > 1:
            doc.add_page_break()
        caption = doc.add_paragraph(f"Segmento {idx} del flujograma formal.")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.bold = True
            run.font.size = Pt(9)
        holder = doc.add_paragraph()
        holder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        holder.add_run().add_picture(str(image), width=Inches(6.35))

    doc.save(str(OUTPUT))
    print(OUTPUT.resolve())


if __name__ == "__main__":
    main()
