from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "entregas" / "entregable-1" / "generated"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_DOCX = OUT_DIR / "Equipo_CriterIA_Entregable_1_StartupEducativa_PREVIEW.docx"

BLUE = RGBColor(31, 78, 121)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GRID = "D9E2F3"
BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(90, 90, 90)


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths=None, font_size=9, header_fill=LIGHT_GRAY):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for idx, cell in enumerate(row.cells):
            if widths:
                set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_run_font(run, size=font_size, color=BLACK)
            if row_idx == 0:
                shade_cell(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, size=font_size, bold=True, color=BLACK)


def add_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, bold=True, color=BLUE)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text="", bold_prefix=None, italic=False, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, bold=True, color=BLACK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11, italic=italic, color=BLACK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, italic=italic, color=BLACK)
    return p


def add_note_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F6F9")
    set_cell_width(cell, 6.3)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=10, italic=True, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Campo"
    hdr[1].text = "Valor"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    style_table(table, widths=[1.65, 4.75], font_size=9.5)
    return table


def add_data_table(doc, headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    style_table(table, widths=widths, font_size=font_size)
    return table


def set_portrait(section):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def set_landscape(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)


def add_document_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.text = "Entregable 1 - Startup Educativa"
    for run in p.runs:
        set_run_font(run, size=9, color=MUTED)


activities = [
    ("1", "Analizar la guía y el contexto", "Todos", "Reunión inicial", "Guía, mapa visual, syllabus", "Guía comprendida y dudas listadas"),
    ("2", "Diseñar el servicio", "Todos", "Sesión de diseño", "Notas de actividad 1", "Dinámica educativa definida y explicable"),
    ("3", "Seleccionar la herramienta de apoyo", "Lucas y Diego", "Trabajo individual + puesta en común", "Acceso a herramientas de IA", "Herramienta elegida para la estación"),
    ("4", "Preparar fichas de candidatos ficticios", "Arnold y Mathias", "Trabajo individual + revisión cruzada", "ADR-0003, plantilla y checklist", "Tres fichas sintéticas aprobadas"),
    ("5", 'Preparar el chatbot "Dos Tonos"', "Lucas y Mathias", "Trabajo de preparación", "Base de candidatos y herramienta seleccionada", "Chatbot listo para demostración"),
    ("6", "Preparar la demostración de la Estación 2", "Todos + sub-equipo B", "Trabajo previo al evento", "Guion y lista de verificación", "Estación lista para operar"),
    ("7", "Prestar el servicio en el colegio", "Todos + sub-equipo B", "Colegio asignado", 'Chatbot "Dos Tonos", dispositivo e internet', "Experiencia realizada y datos registrados"),
    ("8", "Cerrar la estación y ordenar evidencia", "Mathias y Arnold + Módulo 5", "Cierre posterior", "Observaciones, incidencias y datos", "Evidencia ordenada y usable"),
    ("9", "Elaborar el informe académico", "CriterIA completo", "Trabajo final", "Evidencia del servicio y guía", "Documento completo y alineado al servicio"),
]

controls = [
    ("1", "¿Guía comprendida?", "Revisar lectura de la guía y mapa visual"),
    ("2", "¿Objetivo y dinámica son claros?", "Revisar objetivo de aprendizaje"),
    ("3", "¿Herramienta apta para la estación?", "Revisar opciones con acceso real del equipo"),
    ("4", "¿Fichas completas y verificadas?", "Completar o corregir fichas de candidatos"),
    ("5", "¿Chatbot validado?", "Corregir preparación del chatbot"),
    ("6", "¿Estación 2 lista?", "Completar preparación operativa"),
    ("7", "¿Experiencia y datos registrados?", "Registrar incidencia y ajustar evidencia"),
    ("8", "¿Evidencia completa y usable?", "Ordenar observaciones y datos"),
    ("9", "¿Documento explica el servicio?", "Revisar el documento final"),
]

intervenientes = [
    ("Equipo CriterIA", "Diseña el servicio, prepara el chatbot, valida la estación, acompaña la demostración, ordena evidencia y elabora el informe."),
    ("Sub-equipo B de Módulo 2", "Apoya la facilitación, materiales de apoyo y operación de la Estación 2 durante el evento."),
    ("Módulo 5 - Evaluación y Análisis", "Recibe los datos y evidencias generadas durante la experiencia para analizar impacto."),
    ("Docente de la asignatura", "Recibe el informe académico y evalúa el proceso según la guía."),
    ("Estudiantes de educación media", "Participan en la estación, proponen preguntas y comparan dos respuestas del chatbot."),
]

resources = [
    ("Guía del Entregable 1", "Actividad 1", "CriterIA"),
    ("Mapa visual del proyecto Startup Educativa", "Actividad 1", "CriterIA"),
    ("Acceso a herramientas de IA", "Actividades 3 y 5", "Delegados"),
    ("ADR-0003, plantilla y fichas sintéticas", "Actividad 4", "Arnold y Mathias"),
    ('Chatbot "Dos Tonos"', "Actividades 5, 6 y 7", "CriterIA + sub-equipo B"),
    ("Guion de inducción", "Actividad 6", "CriterIA + sub-equipo B"),
    ("Dispositivo e internet", "Actividad 7", "Operador de estación"),
    ("Registro de observaciones e incidencias", "Actividad 8", "CriterIA"),
    ("Datos acordados para evaluación", "Actividad 8", "CriterIA + Módulo 5"),
    ("Formato de entrega PDF", "Actividad 9", "CriterIA"),
]

doc = Document()
set_portrait(doc.sections[0])
add_document_footer(doc.sections[0])

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ENTREGABLE 1")
set_run_font(r, size=24, bold=True, color=BLUE)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Startup Educativa - Estación 2: Alfabetización de Datos")
set_run_font(r, size=14, bold=True, color=BLACK)
p.paragraph_format.space_after = Pt(14)

add_kv_table(doc, [
    ("Asignatura", "Técnicas de Organización y Métodos"),
    ("Equipo", "CriterIA - Sub-equipo A de Módulo 2"),
    ("Servicio", 'Chatbot "Dos Tonos"'),
    ("Destinatario del servicio", "Estudiantes de educación media"),
    ("Destinatario del informe", "Docente de la asignatura"),
    ("Fecha de entrega", "Martes 25 de agosto"),
])

add_note_box(doc, "Nota de alcance: el resultado del servicio no es el PDF. El PDF explica el proceso de diseño, preparación, prestación y cierre de la Estación 2.")

add_heading(doc, "Parte I - Identificación del servicio")
add_para(doc, "Nombre del equipo: CriterIA, sub-equipo A de Módulo 2 (Contenido y Demostraciones), 5 estudiantes universitarios de Ingeniería en Informática.")
add_para(doc, "Tema o componente asignado: Estación 2: Alfabetización de Datos, dentro de los módulos de Diseño de Estaciones y Contenido y Demostraciones.")
add_para(doc, 'Nombre del servicio o experiencia: "Dos Tonos", chatbot de consulta política con dos modos de presentación de una misma base de información.')
add_heading(doc, "Descripción del servicio", level=2)
add_para(doc, '"Dos Tonos" es una experiencia de alfabetización de datos para estudiantes de educación media dentro de la Estación 2 del proyecto Startup Educativa. El estudiante realiza una consulta sobre candidatos políticos y observa dos respuestas generadas a partir de la misma base de información verificada. Una respuesta se presenta en tono neutro y otra con un encuadre intencionalmente diferente, para que el estudiante compare cómo cambia la percepción cuando se resaltan u omiten ciertos aspectos. La actividad no busca orientar el voto, sino enseñar que los datos, aunque sean los mismos, pueden comunicarse de maneras distintas. Durante la estación, el equipo guía la comparación, registra observaciones acordadas con Control de Gestión y deja evidencia del proceso para su evaluación posterior.')

doc.add_page_break()
add_heading(doc, "Parte II - Identificación del proceso")
add_kv_table(doc, [
    ("Nombre", "Diseño, preparación y prestación de la Estación 2: Alfabetización de Datos"),
    ("Inicio", "El equipo CriterIA recibe el componente asignado y define el objetivo educativo de la estación."),
    ("Fin", "Estudiantes de educación media completan la experiencia de comparación y se registran los datos acordados para evaluación."),
    ("Usuario/beneficiario directo", "Estudiantes de educación media que recorren la Estación 2."),
    ("Resultado esperado", "Experiencia educativa clara, guiada y medible sobre distintas formas de presentar una misma información."),
])

add_heading(doc, "Actividades principales", level=2)
add_data_table(
    doc,
    ["N.", "Actividad", "Quién", "Dónde", "Recursos", "Resultado"],
    activities,
    [0.35, 1.65, 1.15, 1.1, 1.25, 1.55],
    font_size=7.8,
)

add_heading(doc, "Parte III - Secuencia del proceso")
add_para(doc, "El proceso se ordena de manera secuencial, con controles al cierre de cada actividad. Si una validación falla, el flujo vuelve al subproceso que debe corregirse.")
add_data_table(
    doc,
    ["Paso", "Actividad", "Control de calidad", "Retorno si NO"],
    [(n, act, controls[int(n)-1][1], controls[int(n)-1][2]) for n, act, *_ in activities],
    [0.45, 2.4, 1.95, 1.75],
    font_size=8.2,
)

doc.add_page_break()
add_heading(doc, "Parte IV - Análisis del proceso")
add_heading(doc, "Intervinientes y funciones", level=2)
add_data_table(doc, ["Interviniente", "Función dentro del proceso"], intervenientes, [1.75, 4.75], font_size=8.8)
add_heading(doc, "Recursos y documentos", level=2)
add_data_table(doc, ["Recurso / documento", "Actividad", "Quién lo utiliza"], resources, [2.6, 1.65, 2.25], font_size=8.5)
add_heading(doc, "Decisiones y controles", level=2)
add_data_table(doc, ["Actividad", "Qué se verifica", "Qué ocurre si NO"], controls, [0.75, 2.65, 3.1], font_size=8.5)

doc.add_page_break()
add_heading(doc, "Parte V - Flujograma del proceso")
add_para(doc, "Tipo de flujograma: vertical, con simbología ANSI. Para que sea legible dentro del Word, el documento incluye una versión compacta en tabla y el flujograma formal detallado dividido por segmentos en anexo horizontal.")
add_heading(doc, "Flujograma compacto para lectura en documento", level=2)
add_data_table(
    doc,
    ["Secuencia", "Actividad principal", "Validación"],
    [(n, act, controls[int(n)-1][1]) for n, act, *_ in activities],
    [0.7, 3.7, 2.1],
    font_size=8.8,
)
add_heading(doc, "Justificación de la elección", level=2)
for item in [
    "La simbología ANSI es reconocible para un lector de Organización y Métodos.",
    "Los rectángulos muestran actividades, los rombos muestran decisiones y las flechas muestran dirección del proceso.",
    "El formato segmentado evita que el flujograma formal quede ilegible al exportar a PDF.",
]:
    p = doc.add_paragraph(style=None)
    p.style = doc.styles["List Bullet"]
    p.add_run(item)

add_heading(doc, "Parte VI - Conclusión del equipo")
add_para(doc, 'Al analizar "Dos Tonos" como un proceso de servicio, el equipo CriterIA aprendió tres cosas centrales. Primero, descomponer la Estación 2 en actividades discretas hace visibles dependencias ocultas: por ejemplo, la demostración depende de que el sub-equipo B coordine la facilitación, de que Módulo 3 confirme las condiciones operativas y de que Módulo 5 defina qué datos necesita recibir. Sin este ejercicio, esas dependencias aparecerían recién el día del evento. Segundo, los puntos de control son tan importantes como las actividades mismas: cada proceso cierra con una validación concreta (guía comprendida, dinámica educativa confirmada, fichas verificadas, estación lista, datos registrados, evidencia ordenada) que evita avanzar con supuestos sin revisar. Tercero, el mismo dato puede producir dos presentaciones muy distintas: la inclinación de los modelos de IA no está en la información que leen, sino en cómo se les pide que la presenten. La verificación "ambos modos leen lo mismo, inclinación solo en Tono A" es la forma operativa de asegurar esta enseñanza.')
add_para(doc, "Los principales aspectos a mejorar antes de la implementación son: definir el alcance del sub-equipo B, construir la base de candidatos, definir la herramienta de apoyo y acordar el contrato de datos con Control de Gestión. El flujograma ayuda a distinguir qué puede paralelizarse y qué depende de una secuencia obligatoria.")

landscape = doc.add_section(WD_SECTION.NEW_PAGE)
set_landscape(landscape)
add_document_footer(landscape)
add_heading(doc, "Anexo A - Flujograma formal detallado")
add_para(doc, "El flujograma formal se divide en segmentos para conservar legibilidad en Word/PDF. La fuente editable permanece como SVG.")

segment_files = [
    ("Segmento 1 - Inicio y Actividad 1", "flujograma-01_inicio_actividad_1.png"),
    ("Segmento 2 - Actividades 2 y 3", "flujograma-02_actividad_2_3.png"),
    ("Segmento 3 - Actividad 4", "flujograma-03_actividad_4.png"),
    ("Segmento 4 - Actividad 5: chatbot", "flujograma-04_actividad_5.png"),
    ("Segmento 5 - Actividad 6", "flujograma-05_actividad_6.png"),
    ("Segmento 6 - Actividad 7: prestación del servicio", "flujograma-06_actividad_7.png"),
    ("Segmento 7 - Actividad 8 y enlace al informe", "flujograma-07_actividad_8.png"),
    ("Segmento 8 - Actividad 9 y cierre del servicio", "flujograma-08_actividad_9_fin.png"),
]

for idx, (caption, filename) in enumerate(segment_files):
    if idx:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    set_run_font(r, size=12, bold=True, color=BLUE)
    img_path = OUT_DIR / filename
    with Image.open(img_path) as image:
        aspect = image.width / image.height
    max_width = 9.2
    max_height = 6.15
    width = min(max_width, max_height * aspect)
    doc.add_picture(str(img_path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(OUT_DOCX)
print(OUT_DOCX)
